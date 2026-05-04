from __future__ import annotations

import re
import textwrap
from collections import defaultdict, deque
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "value_stream_rag_pipeline_full.md"
DOCX_PATH = ROOT / "docs" / "value_stream_rag_pipeline_full.docx"
FALLBACK_DOCX_PATH = ROOT / "docs" / "value_stream_rag_pipeline_full_with_diagrams_fit.docx"
DIAGRAM_DIR = ROOT / "docs" / "value_stream_rag_pipeline_full_diagrams"


EDGE_RE = re.compile(r"^\s*(.+?)\s*-->\s*(?:\|([^|]+)\|\s*)?(.+?)\s*$")
NODE_RE = re.compile(r"^\s*([A-Za-z0-9_]+)(?:\[([^\]]+)\]|\{([^}]+)\})?\s*$")


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            pass
    return ImageFont.load_default()


def parse_node(text: str, labels: dict[str, str], shapes: dict[str, str]) -> str | None:
    match = NODE_RE.match(text.strip())
    if not match:
        return None
    node_id, square_label, diamond_label = match.groups()
    label = square_label or diamond_label
    if label:
        labels[node_id] = label
        shapes[node_id] = "diamond" if diamond_label else "rect"
    else:
        labels.setdefault(node_id, node_id)
        shapes.setdefault(node_id, "rect")
    return node_id


def parse_mermaid(block: str) -> tuple[str, dict[str, str], dict[str, str], list[tuple[str, str, str]]]:
    direction = "TD"
    labels: dict[str, str] = {}
    shapes: dict[str, str] = {}
    edges: list[tuple[str, str, str]] = []

    for raw in block.splitlines():
        line = raw.strip()
        if not line or line.startswith("%%"):
            continue
        if line.startswith("flowchart"):
            parts = line.split()
            if len(parts) > 1:
                direction = parts[1].upper()
            continue
        match = EDGE_RE.match(line)
        if not match:
            continue
        left_raw, edge_label, right_raw = match.groups()
        left = parse_node(left_raw, labels, shapes)
        right = parse_node(right_raw, labels, shapes)
        if left and right:
            edges.append((left, right, edge_label or ""))

    return direction, labels, shapes, edges


def compute_levels(labels: dict[str, str], edges: list[tuple[str, str, str]]) -> dict[str, int]:
    nodes = list(labels)
    outgoing: dict[str, list[str]] = defaultdict(list)
    incoming_count = {node: 0 for node in nodes}
    for src, dst, _ in edges:
        outgoing[src].append(dst)
        incoming_count.setdefault(src, 0)
        incoming_count.setdefault(dst, 0)
        incoming_count[dst] += 1

    queue = deque([node for node, count in incoming_count.items() if count == 0])
    levels = {node: 0 for node in queue}
    while queue:
        node = queue.popleft()
        for child in outgoing.get(node, []):
            levels[child] = max(levels.get(child, 0), levels[node] + 1)
            incoming_count[child] -= 1
            if incoming_count[child] == 0:
                queue.append(child)

    for node in nodes:
        levels.setdefault(node, 0)
    return levels


def wrapped_lines(draw: ImageDraw.ImageDraw, text: str, max_width: int, text_font: ImageFont.ImageFont) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        probe = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), probe, font=text_font)[2] <= max_width:
            current = probe
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [text]


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str) -> None:
    draw.line([start, end], fill=color, width=3)
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) > abs(ey - sy):
        points = [(ex, ey), (ex - 12 if ex > sx else ex + 12, ey - 7), (ex - 12 if ex > sx else ex + 12, ey + 7)]
    else:
        points = [(ex, ey), (ex - 7, ey - 12 if ey > sy else ey + 12), (ex + 7, ey - 12 if ey > sy else ey + 12)]
    draw.polygon(points, fill=color)


def render_mermaid(block: str, output_path: Path, title: str) -> None:
    direction, labels, shapes, edges = parse_mermaid(block)
    levels = compute_levels(labels, edges)
    by_level: dict[int, list[str]] = defaultdict(list)
    for node, level in levels.items():
        by_level[level].append(node)
    for nodes in by_level.values():
        nodes.sort()

    node_w, node_h = 240, 82
    hgap, vgap, margin = 72, 82, 52
    max_level = max(by_level) if by_level else 0
    max_count = max((len(nodes) for nodes in by_level.values()), default=1)

    if direction == "LR":
        width = margin * 2 + (max_level + 1) * node_w + max_level * hgap
        height = margin * 2 + max_count * node_h + max(0, max_count - 1) * vgap
    else:
        width = margin * 2 + max_count * node_w + max(0, max_count - 1) * hgap
        height = margin * 2 + (max_level + 1) * node_h + max_level * vgap
    width = max(width, 900)
    height = max(height, 360)

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = font(22, bold=True)
    node_font = font(17)
    edge_font = font(14)
    draw.text((margin, 18), title, fill="#15313a", font=title_font)

    positions: dict[str, tuple[int, int, int, int]] = {}
    top_offset = 46
    for level, nodes in by_level.items():
        if direction == "LR":
            x = margin + level * (node_w + hgap)
            used_h = len(nodes) * node_h + max(0, len(nodes) - 1) * vgap
            y0 = top_offset + margin + max(0, (height - top_offset - margin * 2 - used_h) // 2)
            for index, node in enumerate(nodes):
                y = y0 + index * (node_h + vgap)
                positions[node] = (x, y, x + node_w, y + node_h)
        else:
            y = top_offset + margin + level * (node_h + vgap)
            used_w = len(nodes) * node_w + max(0, len(nodes) - 1) * hgap
            x0 = margin + max(0, (width - margin * 2 - used_w) // 2)
            for index, node in enumerate(nodes):
                x = x0 + index * (node_w + hgap)
                positions[node] = (x, y, x + node_w, y + node_h)

    for src, dst, edge_label in edges:
        if src not in positions or dst not in positions:
            continue
        sx1, sy1, sx2, sy2 = positions[src]
        dx1, dy1, dx2, dy2 = positions[dst]
        if direction == "LR":
            start = (sx2, (sy1 + sy2) // 2)
            end = (dx1, (dy1 + dy2) // 2)
        else:
            start = ((sx1 + sx2) // 2, sy2)
            end = ((dx1 + dx2) // 2, dy1)
        draw_arrow(draw, start, end, "#64748b")
        if edge_label:
            mx = (start[0] + end[0]) // 2
            my = (start[1] + end[1]) // 2
            bbox = draw.textbbox((0, 0), edge_label, font=edge_font)
            pad = 5
            draw.rounded_rectangle(
                (mx - (bbox[2] - bbox[0]) // 2 - pad, my - 13, mx + (bbox[2] - bbox[0]) // 2 + pad, my + 10),
                radius=6,
                fill="#ffffff",
                outline="#cbd5e1",
            )
            draw.text((mx - (bbox[2] - bbox[0]) // 2, my - 10), edge_label, fill="#334155", font=edge_font)

    for node, box in positions.items():
        x1, y1, x2, y2 = box
        shape = shapes.get(node, "rect")
        fill = "#eefcf8" if shape == "rect" else "#fff7ed"
        outline = "#0f766e" if shape == "rect" else "#c2410c"
        if shape == "diamond":
            cx, cy = (x1 + x2) // 2, (y1 + y2) // 2
            draw.polygon([(cx, y1), (x2, cy), (cx, y2), (x1, cy)], fill=fill, outline=outline)
        else:
            draw.rounded_rectangle(box, radius=10, fill=fill, outline=outline, width=2)
        lines = wrapped_lines(draw, labels[node], node_w - 30, node_font)
        line_h = 20
        text_h = len(lines) * line_h
        y = y1 + (node_h - text_h) // 2
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=node_font)
            draw.text((x1 + (node_w - (bbox[2] - bbox[0])) // 2, y), line, fill="#0f172a", font=node_font)
            y += line_h

    output_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(output_path)


def _row_ink_score(image: Image.Image, y: int) -> int:
    rgb = image.convert("RGB")
    width = rgb.width
    pixels = rgb.crop((0, y, width, min(y + 1, rgb.height))).getdata()
    return sum(1 for r, g, b in pixels if not (r > 245 and g > 245 and b > 245))


def split_tall_image(image_path: Path, max_display_ratio: float = 1.12) -> list[Path]:
    """Split diagrams that would be taller than one Word page at full width."""
    image = Image.open(image_path).convert("RGB")
    max_slice_height = int(image.width * max_display_ratio)
    if image.height <= max_slice_height:
        return [image_path]

    slices: list[Path] = []
    start = 0
    part = 1
    overlap = 28
    while start < image.height:
        target = min(start + max_slice_height, image.height)
        if target < image.height:
            low = max(start + int(max_slice_height * 0.62), target - 220)
            high = min(image.height - 1, target + 220)
            candidates = range(low, high)
            target = min(candidates, key=lambda y: (_row_ink_score(image, y), abs(y - target)))

        crop_start = max(0, start - (overlap if part > 1 else 0))
        crop_end = min(image.height, target + (overlap if target < image.height else 0))
        cropped = image.crop((0, crop_start, image.width, crop_end))
        part_path = image_path.with_name(f"{image_path.stem}_part_{part}{image_path.suffix}")
        cropped.save(part_path)
        slices.append(part_path)
        start = target
        part += 1

    return slices


def add_code(document: Document, lines: list[str]) -> None:
    if not lines:
        return
    paragraph = document.add_paragraph()
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        if index < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def add_table(document: Document, table_lines: list[str]) -> None:
    rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in table_lines]
    if len(rows) >= 2 and is_table_separator(table_lines[1]):
        rows = [rows[0], *rows[2:]]
    if not rows:
        return
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for idx, cell_text in enumerate(rows[0]):
        table.rows[0].cells[idx].text = cell_text
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, cell_text in enumerate(row[: len(cells)]):
            cells[idx].text = cell_text


def build_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    diagram_index = 0
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            language = line.strip().removeprefix("```").strip().lower()
            block: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            if language == "mermaid":
                diagram_index += 1
                image_path = DIAGRAM_DIR / f"diagram_{diagram_index:02d}.png"
                render_mermaid("\n".join(block), image_path, f"Diagram {diagram_index}")
                image_parts = split_tall_image(image_path)
                for part_index, part_path in enumerate(image_parts, start=1):
                    if len(image_parts) > 1:
                        document.add_paragraph(f"Diagram {diagram_index}, part {part_index} of {len(image_parts)}")
                    document.add_picture(str(part_path), width=Inches(6.85))
            else:
                add_code(document, block)
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            table_lines = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(document, table_lines)
            continue

        if not line.strip():
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(heading.group(2), level=level)
            i += 1
            continue

        bullet = re.match(r"^\s*-\s+(.*)$", line)
        if bullet:
            document.add_paragraph(bullet.group(1), style="List Bullet")
            i += 1
            continue

        numbered = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if numbered:
            document.add_paragraph(numbered.group(1), style="List Number")
            i += 1
            continue

        document.add_paragraph(line)
        i += 1

    output_path = DOCX_PATH
    try:
        document.save(output_path)
    except PermissionError:
        output_path = FALLBACK_DOCX_PATH
        document.save(output_path)
        print(f"{DOCX_PATH} is locked; wrote fallback copy instead.")
    print(f"Wrote {output_path}")
    print(f"Wrote {diagram_index} diagram images to {DIAGRAM_DIR}")


if __name__ == "__main__":
    build_docx()
